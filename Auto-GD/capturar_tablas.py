"""
Auto-GD — Fase 2+3: Copiar tabla de Excel y pegar como imagen en Word.

Flujo por cada tabla:
  1. Excel COM: seleccionar rango → CopyPicture al portapapeles
  2. Word  COM: buscar párrafo caption → eliminar imagen placeholder previa
               → pegar imagen del portapapeles justo antes del caption
"""

import os
import re
import shutil
import time
import pythoncom
from pathlib import Path

import win32com.client as win32

# Patrón de entrada de TOC: termina con tabulador + número(s) de página
_TOC_RE = re.compile(r'\t\d+\s*\r?$')

# ─────────────────────────────────────────────────────────────────
# Valores actualmente en la plantilla Word base (MARTINA 1).
# Son los textos que se buscan y se reemplazan con los del proyecto.
# ─────────────────────────────────────────────────────────────────
PLANTILLA_NOMBRE    = "LA MARTINA 1"
PLANTILLA_UBICACION = "Paratebueno, Cundinamarca"  # Ciudad, Departamento (title case)
PLANTILLA_FECHA     = "31/03/2026"                 # DD/MM/YYYY
PLANTILLA_MODULOS   = "1728"                       # número de módulos en la plantilla


# ─────────────────────────────────────────────────────────────────
# CONFIGURACIÓN — ajustar por proyecto
# ─────────────────────────────────────────────────────────────────
EXCEL_CALCULO = (
    r"C:\Users\JesúsAndrésBustilloO\Documents\Auto-GD\output"
    r"\PRGD-CALCULO-LA_MARTINA_1-20260331.xlsm"
)

WORD_BASE = (
    r"C:\Users\JesúsAndrésBustilloO\Documents\GD\Proyectos"
    r"\MARTINA 1 Y 2\MARTINA_1\04. Editables\02. Memoria Retie"
    r"\PRGD-MEMORIA ENEL.docx"
)

WORD_SALIDA = (
    r"C:\Users\JesúsAndrésBustilloO\Documents\Auto-GD\output"
    r"\PRGD-MEMORIA-LA_MARTINA_1-20260331.docx"
)

TABLA_MAP = [
    # Regulación DC — cols B:M, sin fila TOTAL
    {"hoja": "Arreglo mppts SM", "rango": "B2:M32",   "tabla": "reg_dc_inv1", "ancla_word": "regulación DC inversor 1"},
    {"hoja": "Arreglo mppts SM", "rango": "B36:M66",  "tabla": "reg_dc_inv2", "ancla_word": "regulación DC inversor 2"},
    {"hoja": "Arreglo mppts SM", "rango": "B70:M100", "tabla": "reg_dc_inv3", "ancla_word": "regulación DC inversor 3"},
    # Pérdidas DC — cols B:K
    {"hoja": "Pérdidas DC", "rango": "B2:K32",  "tabla": "perd_dc_inv1", "ancla_word": "pérdidas DC inversor 1"},
    {"hoja": "Pérdidas DC", "rango": "B35:K65", "tabla": "perd_dc_inv2", "ancla_word": "pérdidas DC inversor 2"},
    {"hoja": "Pérdidas DC", "rango": "B68:K98", "tabla": "perd_dc_inv3", "ancla_word": "pérdidas DC inversor 3"},
    # AC
    {"hoja": "Regulación AC", "rango": "B2:M11", "tabla": "reg_ac",   "ancla_word": "regulación AC"},
    {"hoja": "Pérdidas AC",   "rango": "B2:I10", "tabla": "perd_ac",  "ancla_word": "pérdidas de energía AC"},
]


# ─────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────
def _es_toc(para) -> bool:
    """
    Devuelve True si el párrafo pertenece a la tabla de contenido.
    Criterios:
      1. El texto crudo termina con tabulador + número de página.
      2. El nombre del estilo contiene 'TOC', 'toc', 'Tabla de contenido' o 'Índice'.
    """
    raw = para.Range.Text
    if _TOC_RE.search(raw):
        return True
    try:
        sn = para.Style.NameLocal.lower()
        if any(tok in sn for tok in ('toc', 'tabla de contenido', 'índice', 'indice')):
            return True
    except Exception:
        pass
    return False


def _buscar_caption(doc, ancla: str) -> int | None:
    """
    Devuelve el índice (1-based) del primer párrafo del CUERPO del documento
    cuyo texto contiene 'ancla' (insensible a mayúsculas).
    Excluye entradas del índice/TOC aunque contengan el mismo texto.
    """
    ancla_low = ancla.lower()
    for i in range(1, doc.Paragraphs.Count + 1):
        para = doc.Paragraphs(i)
        if _es_toc(para):
            continue
        if ancla_low in para.Range.Text.strip().lower():
            return i
    return None


def _eliminar_placeholder(doc, caption_idx: int) -> int:
    """
    Busca InlineShapes en los hasta 3 párrafos ANTES del caption y
    elimina el párrafo completo si contiene una imagen.
    Devuelve el nuevo índice del caption (puede haber bajado en 1).
    """
    for offset in range(1, 4):
        prev_idx = caption_idx - offset
        if prev_idx < 1:
            break
        prev_range = doc.Paragraphs(prev_idx).Range
        p_start = prev_range.Start
        p_end   = prev_range.End

        for j in range(doc.InlineShapes.Count, 0, -1):
            s = doc.InlineShapes(j)
            if p_start <= s.Range.Start and s.Range.End <= p_end + 1:
                # Eliminar el párrafo completo que contiene el placeholder
                prev_range.Delete()
                print(f"      Placeholder eliminado (era {offset} párrafo(s) antes del caption)")
                return caption_idx - 1   # el caption subió un índice

    return caption_idx


def _pegar_antes_caption(wd, doc, caption_idx: int):
    """
    Pega el contenido del portapapeles como imagen justo antes del
    párrafo indicado.  La imagen queda en su propio párrafo.
    """
    caption_start = doc.Paragraphs(caption_idx).Range.Start
    insert_rng = doc.Range(caption_start, caption_start)
    insert_rng.Select()

    # Pegar imagen del portapapeles
    wd.Selection.Paste()

    # Añadir salto de párrafo DESPUÉS de la imagen para separar del caption
    # (el cursor queda al final del contenido pegado)
    wd.Selection.TypeParagraph()


# ─────────────────────────────────────────────────────────────────
# REEMPLAZAR DATOS DEL PROYECTO EN EL WORD (via python-docx)
# Funciona aunque el texto esté partido entre varios runs XML.
# ─────────────────────────────────────────────────────────────────
def _replace_in_para(para, old: str, new: str) -> int:
    """
    Reemplaza `old` por `new` en un párrafo aunque el texto esté
    dividido en varios runs.  Preserva el formato de cada run.
    Devuelve el número de ocurrencias reemplazadas.
    """
    run_texts = [r.text for r in para.runs]
    full = "".join(run_texts)
    if old not in full:
        return 0

    count = 0
    search_from = 0

    while True:
        pos = full.find(old, search_from)
        if pos == -1:
            break

        end = pos + len(old)
        count += 1

        # Construir mapa carácter → índice de run
        run_of = []
        for ri, t in enumerate(run_texts):
            run_of.extend([ri] * len(t))

        r_first = run_of[pos]
        r_last  = run_of[end - 1]

        # Offset dentro del primer run donde empieza la coincidencia
        chars_before_first = sum(len(run_texts[i]) for i in range(r_first))
        off_first = pos - chars_before_first

        # Offset dentro del último run donde termina la coincidencia
        chars_before_last = sum(len(run_texts[i]) for i in range(r_last))
        off_last = end - chars_before_last   # exclusivo

        if r_first == r_last:
            # Todo en un solo run → reemplazo directo
            run_texts[r_first] = (run_texts[r_first][:off_first]
                                  + new
                                  + run_texts[r_first][off_last:])
        else:
            # Primer run: conservar lo anterior + añadir el nuevo texto
            run_texts[r_first] = run_texts[r_first][:off_first] + new
            # Runs intermedios: vaciar
            for ri in range(r_first + 1, r_last):
                run_texts[ri] = ""
            # Último run: descartar la parte que era del texto buscado
            run_texts[r_last] = run_texts[r_last][off_last:]

        # Aplicar al documento
        for ri, r in enumerate(para.runs):
            r.text = run_texts[ri]

        # Recalcular full para la siguiente iteración
        full = "".join(run_texts)
        search_from = pos + len(new)

    return count


def _replace_docx(doc_obj, old: str, new: str) -> int:
    """Reemplaza en todos los párrafos y celdas del documento."""
    from docx.oxml.ns import qn
    total = 0

    # Párrafos del cuerpo
    for para in doc_obj.paragraphs:
        total += _replace_in_para(para, old, new)

    # Celdas de todas las tablas
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += _replace_in_para(para, old, new)

    return total


def reemplazar_datos_proyecto(word_path: str, proyecto: dict):
    """
    Abre el Word con python-docx, aplica todos los reemplazos
    run-a-run y guarda.  Se llama ANTES de abrir con Word COM.
    """
    from docx import Document

    doc = Document(word_path)

    nombre    = proyecto["nombre"].upper()
    ciudad    = proyecto["ciudad"].title()
    depto     = proyecto["departamento"].title()
    ubicacion = f"{ciudad}, {depto}"
    fecha_str = proyecto["fecha"].strftime("%d/%m/%Y")
    modulos   = str(int(proyecto.get("modulos", int(PLANTILLA_MODULOS))))

    reemplazos = [
        (PLANTILLA_NOMBRE,            nombre),
        (PLANTILLA_UBICACION,         ubicacion),
        (PLANTILLA_UBICACION.upper(), ubicacion.upper()),
        (PLANTILLA_FECHA,             fecha_str),
        (PLANTILLA_MODULOS,           modulos),
    ]

    print("  Reemplazando datos del proyecto en el Word:")
    for old, new in reemplazos:
        if old == new:
            continue
        n = _replace_docx(doc, old, new)
        print(f"    «{old}» → «{new}»  ({n} ocurrencia(s))")

    doc.save(word_path)
    print("  Guardado.")


# ─────────────────────────────────────────────────────────────────
# TABLA 27 — NIVEL TENSIÓN DC CORREGIDO POR TEMPERATURA
# ─────────────────────────────────────────────────────────────────
def _fmt(valor: float) -> str:
    """Formatea con 2 decimales y coma decimal (notación colombiana)."""
    return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def actualizar_tabla27(doc, paneles: int, voc_tmin: float, vmpp_tmin: float):
    """
    Busca la tabla cuyo caption siguiente contiene 'nivel tensión DC' o 'tabla 27'
    y actualiza la fila de datos con los valores calculados.
    Columnas: [Paneles en serie | VOC Tmin [V] | Vmpp Tmin [V] | VOC total [V]]
    """
    anclas = ("nivel tensión dc", "nivel tension dc", "tabla 27")
    voc_total = paneles * voc_tmin

    for i in range(1, doc.Tables.Count + 1):
        tbl = doc.Tables(i)
        tbl_end = tbl.Range.End
        # Leer hasta 300 caracteres después de la tabla para encontrar el caption
        check_end = min(tbl_end + 300, doc.Content.End)
        check_text = doc.Range(tbl_end, check_end).Text.strip().lower()

        if any(a in check_text for a in anclas):
            # Verificar que la tabla tenga al menos 2 filas y 4 columnas
            if tbl.Rows.Count < 2 or tbl.Columns.Count < 4:
                continue
            try:
                tbl.Cell(2, 1).Range.Text = str(paneles)
                tbl.Cell(2, 2).Range.Text = _fmt(voc_tmin)
                tbl.Cell(2, 3).Range.Text = _fmt(vmpp_tmin)
                tbl.Cell(2, 4).Range.Text = _fmt(voc_total)
                print(f"  [OK] Tabla 27 actualizada → paneles={paneles}, "
                      f"VOC={_fmt(voc_tmin)}, Vmpp={_fmt(vmpp_tmin)}, VOCtotal={_fmt(voc_total)}")
            except Exception as e:
                print(f"  [ERROR Tabla 27] {e}")
            return

    print("  [ADVERTENCIA] Tabla 27 (Nivel tensión DC) no encontrada en el Word.")


# ─────────────────────────────────────────────────────────────────
# PROCESO PRINCIPAL
# ─────────────────────────────────────────────────────────────────
def ejecutar(excel_path: str, word_base: str, word_salida: str, tabla_map: list,
             paneles_serie: int = 24, proyecto_info: dict = None):
    Path(word_salida).parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(word_base, word_salida)
    print(f"Word base copiado a:\n  {word_salida}\n")

    # ── Reemplazos de texto con python-docx (antes de abrir con COM) ─
    if proyecto_info:
        print(">> Reemplazando datos del proyecto…")
        reemplazar_datos_proyecto(word_salida, proyecto_info)
        print()

    pythoncom.CoInitialize()

    xl = win32.Dispatch("Excel.Application")
    xl.Visible = False
    xl.DisplayAlerts = False

    wd = win32.Dispatch("Word.Application")
    wd.Visible = False
    wd.DisplayAlerts = 0   # wdAlertsNone

    try:
        wb  = xl.Workbooks.Open(os.path.abspath(excel_path))
        xl.Calculate()
        time.sleep(2)

        # Leer valores calculados de la hoja MEMORIA
        ws_mem   = wb.Worksheets("MEMORIA")
        voc_tmin  = ws_mem.Range("E13").Value   # VOC ajustado por Tmin
        vmpp_tmin = ws_mem.Range("E15").Value   # Vmpp ajustado por Tmin
        print(f"MEMORIA E13 (VOC  Tmin) = {voc_tmin}")
        print(f"MEMORIA E15 (Vmpp Tmin) = {vmpp_tmin}\n")

        doc = wd.Documents.Open(os.path.abspath(word_salida))

        for entrada in tabla_map:
            hoja      = entrada["hoja"]
            rango_str = entrada["rango"]
            tabla_id  = entrada["tabla"]
            ancla     = entrada["ancla_word"]

            print(f"[{tabla_id}]")

            # ── 1. Copiar rango como imagen desde Excel ───────────
            try:
                ws = wb.Worksheets(hoja)
                ws.Activate()
                ws.Range(rango_str).CopyPicture(Appearance=1, Format=2)
                time.sleep(0.4)   # dar tiempo al portapapeles
                print(f"  Excel: rango {hoja}!{rango_str} copiado")
            except Exception as e:
                print(f"  [ERROR Excel] {e}")
                continue

            # ── 2. Buscar caption en Word ─────────────────────────
            caption_idx = _buscar_caption(doc, ancla)
            if caption_idx is None:
                print(f"  [ERROR Word] No se encontró: '{ancla}'")
                continue
            print(f"  Word: caption en párrafo {caption_idx} — «{doc.Paragraphs(caption_idx).Range.Text.strip()[:55]}»")

            # ── 3. Eliminar imagen placeholder previa ─────────────
            caption_idx = _eliminar_placeholder(doc, caption_idx)

            # ── 4. Pegar antes del caption ────────────────────────
            try:
                _pegar_antes_caption(wd, doc, caption_idx)
                print(f"  OK: imagen pegada antes del caption")
            except Exception as e:
                print(f"  [ERROR al pegar] {e}")

        # ── Tabla 27: Nivel tensión DC corregido por temperatura ──
        print("\n[tabla27]")
        actualizar_tabla27(doc, paneles_serie, voc_tmin, vmpp_tmin)

        doc.Save()
        doc.Close()
        wb.Close(SaveChanges=False)
        print(f"\n✓ Word guardado en:\n  {word_salida}")

    finally:
        try:
            xl.Quit()
        except Exception:
            pass
        try:
            wd.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


# ─────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=== Auto-GD — Fase 2+3: Excel → portapapeles → Word ===\n")
    ejecutar(EXCEL_CALCULO, WORD_BASE, WORD_SALIDA, TABLA_MAP)
    print("\n=== Proceso completado ===")
